'''import os
import re
from PyPDF2 import PdfReader
from docx import Document
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import nltk
from nltk.corpus import stopwords
from collections import Counter

# Download NLTK data
nltk.download('stopwords')
stop_words = set(stopwords.words('english'))

# Load Spacy model
nlp = spacy.load("en_core_web_sm")

def extract_text_from_file(file_path):
    """Extract text from PDF or DOCX files"""
    if file_path.endswith('.pdf'):
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            text = " ".join([page.extract_text() for page in reader.pages])
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        text = " ".join([para.text for para in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format")
    
    return text

def preprocess_text(text):
    """Clean and preprocess text"""
    # Lowercase
    text = text.lower()
    # Remove special characters
    text = re.sub(r'[^a-zA-Z0-9\s]', '', text)
    # Remove extra whitespace
    text = ' '.join(text.split())
    return text

def extract_keywords(text, n=20):
    """Extract top keywords from text"""
    # Process text with spaCy
    doc = nlp(text)
    
    # Filter out stopwords and punctuation, get lemmas
    keywords = [token.lemma_ for token in doc 
               if not token.is_stop and not token.is_punct and token.is_alpha]
    
    # Count keyword frequencies
    keyword_counts = Counter(keywords)
    
    # Get most common keywords
    top_keywords = [kw for kw, count in keyword_counts.most_common(n)]
    
    return top_keywords

def calculate_similarity(text1, text2):
    """Calculate cosine similarity between two texts"""
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform([text1, text2])
    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
    return similarity * 100  # Convert to percentage

def generate_suggestions(resume_keywords, jd_keywords, missing_keywords):
    """Generate suggestions for improving the resume"""
    suggestions = []
    
    # Convert missing_keywords set to list before slicing
    missing_keywords_list = list(missing_keywords)
    
    # Add missing keywords
    if missing_keywords_list:
        top_keywords = missing_keywords_list[:5]  # Now we can safely slice
        suggestions.append(f"Add these keywords: {', '.join(top_keywords)}")
    
    # Check resume length
    if len(resume_keywords) < 10:
        suggestions.append("Your resume seems too short. Consider adding more relevant content.")
    
    # Check for action verbs
    action_verbs = {'manage', 'lead', 'develop', 'create', 'implement', 'improve', 
                   'analyze', 'design', 'coordinate', 'execute'}
    resume_verbs = set(resume_keywords) & action_verbs
    if len(resume_verbs) < 3:
        suggestions.append("Add more action verbs to make your accomplishments stand out.")
    
    return suggestions

def analyze_resume(resume_path, job_description, roles_responsibilities=""):
    """Main function to analyze resume against job description"""
    # Extract and preprocess texts
    resume_text = extract_text_from_file(resume_path)
    clean_resume = preprocess_text(resume_text)
    clean_jd = preprocess_text(job_description + " " + roles_responsibilities)
    
    # Extract keywords
    resume_keywords = extract_keywords(clean_resume)
    jd_keywords = extract_keywords(clean_jd)
    
    # Find missing keywords
    missing_keywords = set(jd_keywords) - set(resume_keywords)
    
    # Calculate similarity score
    similarity_score = calculate_similarity(clean_resume, clean_jd)
    
    # Generate suggestions
    suggestions = generate_suggestions(resume_keywords, jd_keywords, missing_keywords)
    
    # Prepare results
    results = {
        "score": round(similarity_score, 1),
        "keywords": {
            "required": jd_keywords,
            "present": [kw in resume_keywords for kw in jd_keywords],
            "counts": [resume_text.lower().count(kw) for kw in jd_keywords]
        },
        "suggestions": suggestions
    }
    
    return results'''
import os
import re
from PyPDF2 import PdfReader
from docx import Document
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import nltk
from nltk.corpus import stopwords
from collections import Counter
import textstat
import numpy as np

# Download NLTK data
nltk.download('stopwords')
stop_words = set(stopwords.words('english'))

# Load Spacy model
nlp = spacy.load("en_core_web_sm")

# Action verbs for analysis
ACTION_VERBS = {
    'manage', 'lead', 'develop', 'create', 'implement', 'improve', 'analyze',
    'design', 'coordinate', 'execute', 'optimize', 'build', 'launch', 'increase',
    'reduce', 'enhance', 'innovate', 'transform', 'automate', 'scale'
}

def extract_text_from_file(file_path):
    """Extract text from PDF or DOCX files"""
    try:
        if file_path.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                text = " ".join([page.extract_text() for page in reader.pages if page.extract_text()])
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            text = " ".join([para.text for para in doc.paragraphs if para.text])
        else:
            raise ValueError("Unsupported file format")
        
        return text.strip()
    except Exception as e:
        raise ValueError(f"Error extracting text from file: {str(e)}")

def preprocess_text(text):
    """Clean and preprocess text"""
    # Lowercase
    text = text.lower()
    # Remove special characters
    text = re.sub(r'[^a-zA-Z0-9\s]', '', text)
    # Remove extra whitespace
    text = ' '.join(text.split())
    return text

def calculate_keyword_importance(jd_text):
    """Calculate keyword importance using TF-IDF with enhanced scoring"""
    try:
        # Use both unigrams and bigrams
        vectorizer = TfidfVectorizer(ngram_range=(1, 2), stop_words='english')
        tfidf_matrix = vectorizer.fit_transform([jd_text])
        feature_names = vectorizer.get_feature_names_out()
        importance_scores = np.array(tfidf_matrix.sum(axis=0)).flatten()
        
        # Enhance scores for action verbs and technical terms
        enhanced_scores = {}
        for word, score in zip(feature_names, importance_scores):
            enhanced_score = score
            # Boost action verbs
            if word in ACTION_VERBS:
                enhanced_score *= 1.5
            # Boost technical terms (words in all caps or with numbers)
            if re.search(r'[A-Z]{2,}|\d', word):
                enhanced_score *= 1.3
            enhanced_scores[word] = enhanced_score
        
        return enhanced_scores
    except:
        # Fallback to simple word count if TF-IDF fails
        words = re.findall(r'\b\w+\b', jd_text.lower())
        return Counter(words)

def extract_keywords(text, n=30):
    """Extract top keywords from text with importance"""
    try:
        # Process text with spaCy
        doc = nlp(text)
        
        # Filter out stopwords and punctuation, get lemmas
        keywords = [token.lemma_ for token in doc 
                   if not token.is_stop and not token.is_punct and token.is_alpha]
        
        # Count keyword frequencies
        keyword_counts = Counter(keywords)
        
        # Get most common keywords
        top_keywords = [kw for kw, count in keyword_counts.most_common(n)]
        
        return top_keywords
    except Exception as e:
        print(f"Error extracting keywords: {str(e)}")
        return []

def analyze_content_metrics(resume_text, jd_text):
    """Calculate various content metrics with enhanced analysis"""
    metrics = {}
    
    try:
        # Count action verbs (enhanced detection)
        doc = nlp(resume_text.lower())
        action_verbs_count = sum(
            1 for token in doc 
            if token.lemma_ in ACTION_VERBS and token.pos_ == 'VERB'
        )
        metrics['action_verbs'] = min(100, action_verbs_count * 8)  # Scale to 100
        
        # Count quantified achievements (enhanced pattern matching)
        quant_patterns = [
            r'\b\d+%', r'\$\d+', r'\d+\+', r'\d+\s*(years|months)', 
            r'increased by', r'reduced by', r'saved\s*\$\d+',
            r'improved\s*by', r'optimized\s*by'
        ]
        quantified = sum(
            len(re.findall(pattern, resume_text, re.IGNORECASE))
            for pattern in quant_patterns
        )
        metrics['quantified_achievements'] = min(100, quantified * 15)  # Scale to 100
        
        # Skills match percentage
        jd_skills = set(re.findall(r'\b[A-Z][a-zA-Z]+\b', jd_text))  # Capitalized words
        resume_skills = set(re.findall(r'\b[A-Z][a-zA-Z]+\b', resume_text))
        metrics['skills_match'] = (
            len(jd_skills & resume_skills) / len(jd_skills) * 100
            if jd_skills else 0
        )
        
        # Section completeness (enhanced check)
        sections = ['experience', 'education', 'skills', 'projects']
        found_sections = sum(
            1 for section in sections 
            if re.search(rf'\b{section}\b', resume_text.lower())
        )
        metrics['section_completeness'] = (found_sections / len(sections)) * 100
        
        return metrics
    except Exception as e:
        print(f"Error analyzing content metrics: {str(e)}")
        return {
            'action_verbs': 0,
            'quantified_achievements': 0,
            'skills_match': 0,
            'section_completeness': 0
        }

def generate_suggestions(resume_text, jd_text, missing_keywords, metrics):
    """Generate suggestions for improving the resume"""
    suggestions = {
        'Keywords': [],
        'Content': [],
        'Style': []
    }
    
    try:
        # Convert missing_keywords set to list before slicing
        missing_keywords_list = list(missing_keywords)
        
        # Keyword suggestions
        if missing_keywords_list:
            top_keywords = missing_keywords_list[:5]
            suggestions['Keywords'].append(f"Add these important keywords: {', '.join(top_keywords)}")
        
        if len(missing_keywords_list) > 5:
            suggestions['Keywords'].append(f"Plus {len(missing_keywords_list) - 5} more important keywords to include")
        
        # Content suggestions
        if metrics['action_verbs'] < 50:
            suggestions['Content'].append("Add more action verbs to make your accomplishments stand out")
        
        if metrics['quantified_achievements'] < 50:
            suggestions['Content'].append("Include more quantified achievements (e.g., 'Increased sales by 30%')")
        
        if metrics['skills_match'] < 70:
            suggestions['Content'].append("Highlight more of the required skills from the job description")
        
        if metrics['section_completeness'] < 100:
            suggestions['Content'].append("Ensure all major sections (Experience, Education, Skills, Projects) are included")
        
        # Style suggestions
        readability = textstat.flesch_reading_ease(resume_text)
        if readability < 60:
            suggestions['Style'].append("Simplify some sentences to improve readability (aim for 8th grade reading level)")
        
        # Check resume length
        word_count = len(resume_text.split())
        if word_count < 200:
            suggestions['Style'].append("Your resume seems too short. Consider adding more relevant content")
        elif word_count > 800:
            suggestions['Style'].append("Your resume might be too long. Consider being more concise")
        
        return suggestions
    except Exception as e:
        print(f"Error generating suggestions: {str(e)}")
        return suggestions

def analyze_resume(resume_path, job_description):
    """Advanced resume analysis against job description"""
    try:
        # Extract and preprocess texts
        resume_text = extract_text_from_file(resume_path)
        clean_resume = preprocess_text(resume_text)
        clean_jd = preprocess_text(job_description)
        
        # Calculate keyword importance from JD
        keyword_importance = calculate_keyword_importance(clean_jd)
        
        # Extract keywords
        resume_keywords = extract_keywords(clean_resume)
        jd_keywords = extract_keywords(clean_jd)
        
        # Find missing keywords
        missing_keywords = set(jd_keywords) - set(resume_keywords)
        
        # Calculate scores
        vectorizer = TfidfVectorizer()
        tfidf_matrix = vectorizer.fit_transform([clean_resume, clean_jd])
        similarity_score = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0] * 100
        
        # Calculate keyword match percentage
        keyword_match = (len(set(resume_keywords) & set(jd_keywords)) / len(set(jd_keywords))) * 100
        
        # Analyze content metrics
        metrics = analyze_content_metrics(resume_text, clean_jd)
        
        # Generate suggestions
        suggestions = generate_suggestions(resume_text, clean_jd, missing_keywords, metrics)
        
        # Prepare keyword analysis data
        keyword_data = {
            "required": jd_keywords,
            "present": [kw in resume_keywords for kw in jd_keywords],
            "counts": [resume_text.lower().count(kw) for kw in jd_keywords],
            "importance": [keyword_importance.get(kw, 0) for kw in jd_keywords]
        }
        
        # Prepare results
        results = {
            "score": round(similarity_score, 1),
            "keyword_match": round(keyword_match, 1),
            "content_relevance": round(np.mean(list(metrics.values())), 1),
            "keywords": keyword_data,
            "metrics": metrics,
            "suggestions": suggestions
        }
        
        return results
    except Exception as e:
        print(f"Error analyzing resume: {str(e)}")
        raise e